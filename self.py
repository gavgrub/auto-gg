from json import load
from datetime import date
from docx import Document
from docx.shared import Inches

from helper.ai import generateText
from helper.doc import formatDoc
from helper.text import ordinal

# Load Data for Document
with open("data.json", "r", encoding="utf-8") as f:
    data = load(f)

today = date.today()

# Create document
doc = Document()

# Set margins
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

# Header
doc.add_paragraph().add_run("Student Experiential Learning Fund (SELF)").bold = True
doc.add_paragraph("Application I")
doc.add_paragraph("")
doc.add_paragraph("Schulich School of Engineering")
doc.add_paragraph("University of Calgary")
doc.add_paragraph("")
doc.add_paragraph(f"Revised {today.strftime('%B')} {today.day}{ordinal(today.day)}, {today.year}")
doc.add_paragraph("")

# Primary Applicant Information
doc.add_paragraph().add_run("Primary Applicant Information").bold = True

table = doc.add_table(rows=6, cols=2)

table.cell(0, 0).text = "First Name"
table.cell(0, 1).text = data["firstName"]

table.cell(1, 0).text = "Last Name"
table.cell(1, 1).text = data["lastName"]

table.cell(2, 0).text = "UCalgary Email"
table.cell(2, 1).text = data["email"]

table.cell(3, 0).text = "UCID"
table.cell(3, 1).text = data["ucid"]

table.cell(4, 0).text = "Club / Team Name"
table.cell(4, 1).text = data["club"]

table.cell(5, 0).text = "Position on Club / Team"
table.cell(5, 1).text = data["position"]

doc.add_paragraph("")

# Activity Proposal
doc.add_paragraph().add_run("Activity Proposal").bold = True

table = doc.add_table(rows=9, cols=2)

table.cell(0, 0).text = "Type of Activity (SELF Category)"
table.cell(0, 1).text = data["activityType"]

table.cell(1, 0).text = "Activity Name"
table.cell(1, 1).text = data["activityName"]

table.cell(2, 0).merge(table.cell(2, 1))
table.cell(2, 0).text = "Description of Activity"

table.cell(3, 0).merge(table.cell(3, 1))
table.cell(3, 0).text = "TO ADD" # ADD IN THE AI

table.cell(4, 0).text = "Start Date of Activity"
table.cell(4, 1).text = data["startDate"]

table.cell(5, 0).text = "End Date of Activity"
table.cell(5, 1).text = data["endDate"]

table.cell(6, 0).text = "Location of Activity"
table.cell(6, 1).text = data["location"]

table.cell(7, 0).merge(table.cell(7, 1))
table.cell(7, 0).text = f"Schedule / Itinerary of Activity ({data['time']})"

table.cell(8, 0).merge(table.cell(8, 1))
table.cell(8, 0).text = f"TO ADD" # ADD IN THE AI

doc.add_paragraph("")

# Reflection
doc.add_paragraph().add_run("Reflection").bold = True

table = doc.add_table(rows=4, cols=1)

table.cell(0, 0).text = "How does this activity qualify under this category?"
table.cell(1, 0).text = "TO ADD" # ADD IN THE AI
table.cell(2, 0).text = "How does this activity have an impact on Experiential Learning and/or have relevance towards the betterment of the Engineering student experience?"
table.cell(3, 0).text = "TO ADD" # ADD IN THE AI

doc.add_paragraph("")

# Budget
doc.add_paragraph().add_run("Budget").bold = True
doc.add_paragraph("Please refer to \"SELF Budget\" excel sheet.")

# Format the document correctly
formatDoc(doc)

# Save the document
doc.save(f"{data['activityName']} SELF Application.docx")
