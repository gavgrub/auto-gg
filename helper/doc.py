from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

# Length -> twips (handles int / Length)
def tw(x):
    return x.twips if hasattr(x, "twips") else int(x / 635)

# Format entire document
def formatDoc(doc):
    sec = doc.sections[0]

    for p in doc.paragraphs:
        pf = p.paragraph_format
        pf.line_spacing = 1.15
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

        for r in p.runs:
            r.font.name = "Century"
            r.font.size = Pt(12)

    for t in doc.tables:
        formatTable(t, sec)

# Format one table to fit margins
def formatTable(table, section, label_ratio=0.35, style="Table Grid"):
    usable = section.page_width - section.left_margin - section.right_margin

    table.style = style

    table.autofit = False
    _setTblWidth(table, usable)
    _removeIndent(table)

    lw = usable * label_ratio
    vw = usable * (1 - label_ratio)

    for row in table.rows:
        cells = row.cells

        if (len(cells) == 2):
            cells[0].width = lw
            cells[1].width = vw
        else:
            cells[0].width = usable

        for cell in row.cells:
            _cellMargins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.name = "Century"
                    r.font.size = Pt(12)

# Force table width
def _setTblWidth(table, width):
    tblPr = table._tbl.tblPr
    for x in tblPr.findall(qn("w:tblW")):
        tblPr.remove(x)

    w = OxmlElement("w:tblW")
    w.set(qn("w:w"), str(tw(width)))
    w.set(qn("w:type"), "dxa")
    tblPr.append(w)

# Remove Word indentation
def _removeIndent(table):
    tblPr = table._tbl.tblPr
    for x in tblPr.findall(qn("w:tblInd")):
        tblPr.remove(x)

    ind = OxmlElement("w:tblInd")
    ind.set(qn("w:w"), "0")
    ind.set(qn("w:type"), "dxa")
    tblPr.append(ind)

# Cell padding
def _cellMargins(cell, m=144):
    tcPr = cell._tc.get_or_add_tcPr()
    for x in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(x)

    mar = OxmlElement("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        n = OxmlElement(f"w:{side}")
        n.set(qn("w:w"), str(m))
        n.set(qn("w:type"), "dxa")
        mar.append(n)

    tcPr.append(mar)